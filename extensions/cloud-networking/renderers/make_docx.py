#!/usr/bin/env python3
"""make_docx.py - Cloud Networking Word renderer (with REAL Word styles).

Install: pip install python-docx markdown-it-py
Usage:   python make_docx.py --input report.md --output report.docx --specialist "Hybrid Connectivity"

Policy mandates real Word styles (Heading 1/2/3, Title, Quote, List Bullet) -
never hand-roll bold runs to fake structure. A TOC field is inserted at the
top so Word can populate the table of contents when opened.
"""
from __future__ import annotations
import argparse, datetime, html, pathlib, re, sys


PRIMARY_RGB = (0x0E, 0x5A, 0x9C)

_ENTITY_RE = re.compile(r"\\?&(#\d+|#x[0-9a-fA-F]+|[a-zA-Z][a-zA-Z0-9]{1,31});")


def decode_entities(text: str) -> str:
    """Normalize HTML entity references in markdown source to Unicode chars.

    Keeps DOCX in parity with PDF/HTML. markdown-it-py's `inline.content` field
    preserves raw entity text (`&mdash;`), which this renderer writes straight
    into the .text property of table cells - so without this pre-pass the entity
    name shows literally in Word.
    """
    cur = text
    for _ in range(3):
        new = _ENTITY_RE.sub(lambda m: html.unescape("&" + m.group(1) + ";"), cur)
        if new == cur:
            break
        cur = new
    return cur


def add_toc_field(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def style_headings(doc):
    from docx.shared import RGBColor, Pt
    for level in (1, 2, 3, "Title"):
        name = f"Heading {level}" if isinstance(level, int) else level
        try:
            s = doc.styles[name]
            s.font.color.rgb = RGBColor(*PRIMARY_RGB)
        except KeyError:
            pass
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(11)


def render_markdown(doc, md_text: str) -> int:
    """Render markdown into the docx using real styles. Returns paragraph count."""
    try:
        from markdown_it import MarkdownIt
    except ImportError:
        sys.exit("ERROR: markdown-it-py not installed. Run: pip install markdown-it-py")

    md = MarkdownIt("commonmark", {"breaks": False, "html": False}).enable("table").enable("strikethrough")
    tokens = md.parse(md_text)
    para_count = 0
    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "heading_open":
            level = int(tok.tag[1])
            content_tok = tokens[i + 1]
            style = f"Heading {min(level, 9)}"
            try:
                doc.add_paragraph(content_tok.content, style=style)
            except KeyError:
                doc.add_paragraph(content_tok.content, style="Heading 1")
            para_count += 1
            i += 3
            continue
        if tok.type == "paragraph_open":
            content_tok = tokens[i + 1]
            text = content_tok.content
            doc.add_paragraph(text)
            para_count += 1
            i += 3
            continue
        if tok.type == "bullet_list_open":
            j = i + 1
            while j < len(tokens) and tokens[j].type != "bullet_list_close":
                if tokens[j].type == "inline":
                    try:
                        doc.add_paragraph(tokens[j].content, style="List Bullet")
                    except KeyError:
                        doc.add_paragraph(f"- {tokens[j].content}")
                    para_count += 1
                j += 1
            i = j + 1
            continue
        if tok.type == "ordered_list_open":
            j = i + 1
            n = 1
            while j < len(tokens) and tokens[j].type != "ordered_list_close":
                if tokens[j].type == "inline":
                    try:
                        doc.add_paragraph(tokens[j].content, style="List Number")
                    except KeyError:
                        doc.add_paragraph(f"{n}. {tokens[j].content}")
                    n += 1
                    para_count += 1
                j += 1
            i = j + 1
            continue
        if tok.type == "blockquote_open":
            j = i + 1
            while j < len(tokens) and tokens[j].type != "blockquote_close":
                if tokens[j].type == "inline":
                    try:
                        doc.add_paragraph(tokens[j].content, style="Quote")
                    except KeyError:
                        doc.add_paragraph(f"> {tokens[j].content}")
                    para_count += 1
                j += 1
            i = j + 1
            continue
        if tok.type == "table_open":
            rows = []
            current_row = []
            j = i + 1
            while j < len(tokens) and tokens[j].type != "table_close":
                if tokens[j].type == "inline":
                    current_row.append(tokens[j].content)
                if tokens[j].type in ("tr_close",):
                    if current_row:
                        rows.append(current_row)
                        current_row = []
                j += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                try:
                    table.style = "Light Grid Accent 1"
                except KeyError:
                    pass
                for r_idx, r in enumerate(rows):
                    for c_idx, cell_val in enumerate(r):
                        table.cell(r_idx, c_idx).text = cell_val
                para_count += len(rows)
            i = j + 1
            continue
        i += 1
    return para_count


def main() -> int:
    ap = argparse.ArgumentParser(description="Render a Cloud Networking markdown report to DOCX with real Word styles.")
    ap.add_argument("--input", required=True, type=pathlib.Path)
    ap.add_argument("--output", required=True, type=pathlib.Path)
    ap.add_argument("--specialist", default="Cloud Networking")
    args = ap.parse_args()

    try:
        from docx import Document
    except ImportError:
        sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")

    md_text = args.input.read_text(encoding="utf-8")
    md_text = decode_entities(md_text)
    doc = Document()
    style_headings(doc)

    doc.add_paragraph(args.specialist, style="Title")
    today = datetime.date.today().isoformat()
    doc.add_paragraph(f"Cloud Networking - {today}")

    add_toc_field(doc)
    doc.add_page_break()

    para_count = render_markdown(doc, md_text)

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.text = f"Cloud Networking - {args.specialist} - {today}"

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)

    size = args.output.stat().st_size
    print(f"OK  {args.output}  ({size:,} bytes, ~{para_count} paragraphs)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
